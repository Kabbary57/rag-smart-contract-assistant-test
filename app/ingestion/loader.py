"""
loader.py – Load PDF and DOCX files into LangChain Document objects.

Supports:
  - PDF via PyMuPDF (fitz) with fallback to pdfplumber
  - DOCX via python-docx
  - Plain text files
"""

import logging
import os
from pathlib import Path
from typing import List

from langchain_core.documents import Document

logger = logging.getLogger(__name__)


# ─── PDF Loading ────────────────────────────────────────────────────────────

def _load_pdf_pymupdf(file_path: str) -> List[Document]:
    """Load PDF using PyMuPDF (faster, handles most PDFs)."""
    import fitz  # PyMuPDF

    docs = []
    pdf = fitz.open(file_path)
    file_name = Path(file_path).name

    for page_num in range(len(pdf)):
        page = pdf[page_num]
        text = page.get_text("text").strip()
        if text:
            docs.append(Document(
                page_content=text,
                metadata={
                    "source": file_name,
                    "page": page_num + 1,
                    "total_pages": len(pdf),
                    "file_type": "pdf",
                }
            ))
    pdf.close()
    logger.info(f"PyMuPDF loaded {len(docs)} pages from {file_name}")
    return docs


def _load_pdf_pdfplumber(file_path: str) -> List[Document]:
    """Fallback PDF loader using pdfplumber (better for tables)."""
    import pdfplumber

    docs = []
    file_name = Path(file_path).name

    with pdfplumber.open(file_path) as pdf:
        total_pages = len(pdf.pages)
        for page_num, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                docs.append(Document(
                    page_content=text,
                    metadata={
                        "source": file_name,
                        "page": page_num + 1,
                        "total_pages": total_pages,
                        "file_type": "pdf",
                    }
                ))
    logger.info(f"pdfplumber loaded {len(docs)} pages from {file_name}")
    return docs


def load_pdf(file_path: str) -> List[Document]:
    """Load a PDF file, trying PyMuPDF first then pdfplumber."""
    try:
        docs = _load_pdf_pymupdf(file_path)
        if docs:
            return docs
    except ImportError:
        logger.warning("PyMuPDF not available, falling back to pdfplumber.")
    except Exception as e:
        logger.warning(f"PyMuPDF failed ({e}), falling back to pdfplumber.")

    return _load_pdf_pdfplumber(file_path)


# ─── DOCX Loading ───────────────────────────────────────────────────────────

def load_docx(file_path: str) -> List[Document]:
    """Load a DOCX file paragraph-by-paragraph."""
    from docx import Document as DocxDocument

    file_name = Path(file_path).name
    doc = DocxDocument(file_path)
    full_text = "\n".join(
        para.text for para in doc.paragraphs if para.text.strip()
    )

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                full_text += "\n" + row_text

    docs = [Document(
        page_content=full_text,
        metadata={
            "source": file_name,
            "page": 1,
            "total_pages": 1,
            "file_type": "docx",
        }
    )]
    logger.info(f"python-docx loaded {len(full_text)} chars from {file_name}")
    return docs


# ─── Plain Text ─────────────────────────────────────────────────────────────

def load_txt(file_path: str) -> List[Document]:
    """Load a plain text file."""
    file_name = Path(file_path).name
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()

    return [Document(
        page_content=text,
        metadata={
            "source": file_name,
            "page": 1,
            "total_pages": 1,
            "file_type": "txt",
        }
    )]


# ─── Unified Loader ─────────────────────────────────────────────────────────

def load_document(file_path: str) -> List[Document]:
    """
    Unified entry point.  Dispatches based on file extension.

    Args:
        file_path: Absolute or relative path to the document.

    Returns:
        List of LangChain Document objects (one per page/section).

    Raises:
        ValueError: If the file type is not supported.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = path.suffix.lower()
    size_mb = path.stat().st_size / (1024 * 1024)
    logger.info(f"Loading {path.name} ({size_mb:.1f} MB, type={ext})")

    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return load_docx(file_path)
    elif ext in (".txt", ".md"):
        return load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx, .txt, .md")
